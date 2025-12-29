#!/usr/bin/env python3
"""
Generate academic website and CV from a single papers data file.
Usage: python3 generate_academic_materials.py
"""

import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def load_papers():
    """Load papers from JSON file."""
    with open('papers.json', 'r') as f:
        return json.load(f)['papers']

def generate_website(papers):
    """Generate index.html from papers data."""
    html = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Swaminathan Balasubramaniam</title>
    <style>
        body {
            font-family: 'Georgia', serif;
            line-height: 1.6;
            max-width: 900px;
            margin: 0 auto;
            padding: 40px 20px;
            color: #333;
        }
        
        h1 {
            color: #2c3e50;
            margin-bottom: 10px;
            font-size: 2.2em;
        }
        
        h2 {
            color: #34495e;
            border-bottom: 2px solid #3498db;
            padding-bottom: 10px;
            margin-top: 40px;
            margin-bottom: 20px;
        }
        
        .intro {
            margin-bottom: 30px;
            line-height: 1.8;
        }
        
        .cv-link {
            color: #3498db;
            text-decoration: none;
        }
        
        .cv-link:hover {
            text-decoration: underline;
        }
        
        .paper {
            margin-bottom: 30px;
            padding-bottom: 20px;
            border-bottom: 1px solid #eee;
        }
        
        .paper:last-child {
            border-bottom: none;
        }
        
        .paper-title {
            font-weight: bold;
            font-size: 1.1em;
            color: #2c3e50;
            margin-bottom: 5px;
        }
        
        .paper-authors {
            color: #555;
            margin-bottom: 5px;
            font-style: italic;
        }
        
        .paper-status {
            color: #27ae60;
            font-weight: bold;
            margin-bottom: 5px;
        }
        
        .paper-links a {
            color: #3498db;
            text-decoration: none;
            margin-right: 15px;
        }
        
        .paper-links a:hover {
            text-decoration: underline;
        }
        
        .contact {
            margin-top: 40px;
            padding-top: 20px;
            border-top: 2px solid #3498db;
        }
    </style>
</head>
<body>
    <h1>Swaminathan Balasubramaniam</h1>
    
    <div class="intro">
        <p>I am an Assistant Professor of Finance at NEOMA Business School, France (Paris, Reims, Rouen). <a href="CV_Swaminathan_Balasubramaniam.pdf" class="cv-link" target="_blank">[CV]</a></p>
        
        <p>I'm a financial economist with research interests broadly focused on financial markets and financial institutions. I study challenges faced by financial systems in aggregating information, providing liquidity and preventing coordination failures.</p>
        
        <p>I teach Fintech and Decentralized Finance in the Autumn semester and Blockchain and Fintech in the Spring semester.</p>
    </div>
    
    <h2>Research</h2>
"""
    
    # Add papers
    for paper in papers:
        html += '    <div class="paper">\n'
        html += f'        <div class="paper-title">{paper["title"]}</div>\n'
        
        if paper.get('authors'):
            authors_str = "with " + " and ".join(paper['authors'])
            html += f'        <div class="paper-authors">{authors_str}</div>\n'
        
        if paper.get('status'):
            html += f'        <div class="paper-status">{paper["status"]}</div>\n'
        
        html += '        <div class="paper-links">\n'
        for link_type, url in paper['links'].items():
            label = link_type.upper() if link_type == 'ssrn' else link_type.replace('_', ' ').title()
            html += f'            <a href="{url}" target="_blank">{label}</a>\n'
        html += '        </div>\n'
        html += '    </div>\n    \n'
    
    html += """    <div class="contact">
        <h2>Contact</h2>
        <p>Email: s.balasubramaniam@neoma-bs.fr</p>
    </div>
</body>
</html>
"""
    
    with open('index.html', 'w') as f:
        f.write(html)
    print("✓ Generated index.html")

def generate_cv(papers):
    """Generate CV from papers data."""
    doc = Document()
    
    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Header with name
    name = doc.add_paragraph()
    name_run = name.add_run('SWAMINATHAN BALASUBRAMANIAM')
    name_run.font.size = Pt(16)
    name_run.font.bold = True
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact info
    contact = doc.add_paragraph()
    contact_run = contact.add_run('156 rue de Tolbiac, Paris - 75013\nPhone: +33 678899943 | E-mail: s.balasubramaniam@neoma-bs.fr')
    contact_run.font.size = Pt(10)
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # spacing
    
    # Academic Employment
    heading = doc.add_paragraph()
    heading_run = heading.add_run('ACADEMIC EMPLOYMENT')
    heading_run.font.size = Pt(12)
    heading_run.font.bold = True
    heading_run.font.color.rgb = RGBColor(0, 0, 128)
    
    job = doc.add_paragraph(style='List Bullet')
    job_run = job.add_run('Jan 2022 onwards: NEOMA Business School, Assistant Professor of Finance')
    job_run.font.size = Pt(11)
    
    # Education
    heading = doc.add_paragraph()
    heading_run = heading.add_run('EDUCATION')
    heading_run.font.size = Pt(12)
    heading_run.font.bold = True
    heading_run.font.color.rgb = RGBColor(0, 0, 128)
    
    edu1 = doc.add_paragraph(style='List Bullet')
    edu1_run = edu1.add_run('2015-2021: Washington University in St. Louis, USA - PhD, Finance (Olin Business School)')
    edu1_run.font.size = Pt(11)
    
    edu2 = doc.add_paragraph(style='List Bullet')
    edu2_run = edu2.add_run('2004-06: Indian Institute of Management, Lucknow - MBA with specialization in Finance')
    edu2_run.font.size = Pt(11)
    
    edu3 = doc.add_paragraph(style='List Bullet')
    edu3_run = edu3.add_run('2000-04: Indian Institute of Technology, Kharagpur - B.Tech (Hons.) Biotechnology and Biochemical Engineering')
    edu3_run.font.size = Pt(11)
    
    # Research Interests
    heading = doc.add_paragraph()
    heading_run = heading.add_run('AREAS OF INTEREST')
    heading_run.font.size = Pt(12)
    heading_run.font.bold = True
    heading_run.font.color.rgb = RGBColor(0, 0, 128)
    
    interests = doc.add_paragraph()
    interests_run = interests.add_run('Financial markets and financial institutions (Information aggregation, liquidity, coordination failures)')
    interests_run.font.size = Pt(11)
    
    # Working Papers
    heading = doc.add_paragraph()
    heading_run = heading.add_run('RESEARCH')
    heading_run.font.size = Pt(12)
    heading_run.font.bold = True
    heading_run.font.color.rgb = RGBColor(0, 0, 128)
    
    for paper in papers:
        paper_para = doc.add_paragraph(style='List Bullet')
        
        # Title
        title_run = paper_para.add_run(paper['title'])
        title_run.font.size = Pt(11)
        title_run.font.italic = True
        
        # Authors
        if paper.get('authors'):
            authors_str = " (with " + " and ".join(paper['authors']) + ")"
            authors_run = paper_para.add_run(authors_str)
            authors_run.font.size = Pt(11)
        
        # Status
        if paper.get('status'):
            paper_para.add_run('\n')
            status_run = paper_para.add_run(paper['status'])
            status_run.font.size = Pt(10)
            status_run.font.color.rgb = RGBColor(0, 128, 0)
    
    # Teaching
    heading = doc.add_paragraph()
    heading_run = heading.add_run('TEACHING EXPERIENCE')
    heading_run.font.size = Pt(12)
    heading_run.font.bold = True
    heading_run.font.color.rgb = RGBColor(0, 0, 128)
    
    teach1 = doc.add_paragraph(style='List Bullet')
    teach1_run = teach1.add_run('Fintech and Decentralized Finance (NEOMA Business School)')
    teach1_run.font.size = Pt(11)
    
    teach2 = doc.add_paragraph(style='List Bullet')
    teach2_run = teach2.add_run('Blockchain and Fintech (NEOMA Business School)')
    teach2_run.font.size = Pt(11)
    
    teach3 = doc.add_paragraph(style='List Bullet')
    teach3_run = teach3.add_run('Financial Decisions Under Uncertainty (NEOMA Business School)')
    teach3_run.font.size = Pt(11)
    
    # Industry Experience
    heading = doc.add_paragraph()
    heading_run = heading.add_run('INDUSTRY EXPERIENCE')
    heading_run.font.size = Pt(12)
    heading_run.font.bold = True
    heading_run.font.color.rgb = RGBColor(0, 0, 128)
    
    industry = doc.add_paragraph(style='List Bullet')
    industry_run = industry.add_run('2006-14: Vice President, Axis Capital, Mumbai, India (Investment Banking - Financial Institutions group). Equity raising/M&A in banking sector. Previously: Associate, SBI Caps.')
    industry_run.font.size = Pt(11)
    
    # Save
    doc.save('CV_Swaminathan_Balasubramaniam.docx')
    print("✓ Generated CV_Swaminathan_Balasubramaniam.docx")

def main():
    papers = load_papers()
    generate_website(papers)
    generate_cv(papers)
    print("\n✓ All files generated successfully!")
    print("\nTo update in the future:")
    print("1. Edit papers.json to add/modify papers")
    print("2. Run: python3 generate_academic_materials.py")
    print("3. Upload both index.html and the new CV PDF to GitHub")

if __name__ == '__main__':
    main()
